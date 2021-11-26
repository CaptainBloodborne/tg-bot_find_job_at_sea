# coding=utf-8
"""
main.py - body of Telegram-bot
"""
import telebot
from docx import Document
from telebot import types

from const import TOKEN

bot = telebot.TeleBot(TOKEN)


@bot.message_handler(commands=["start"])
def startmsg(message):
    """

    :param message:
    :return:
    """

    startmenu = types.ReplyKeyboardMarkup(True, False)
    startmenu.row("Get underway!")
    bot.send_message(
        message.chat.id,
        text="""Hooray! Welcome on board!🚢 \nJohn Silver will help you to find a job and get 💰""",
        reply_markup=startmenu,
    )


@bot.message_handler(commands=["info"])
def infomsg(message):
    """

    :param message:
    :return:
    """
    bot.send_message(message.chat.id, text="""My bot""")


@bot.message_handler(content_types=["text"])
def main_handler(message):
    """

    :param message:
    :return:
    """
    if message.text == "Get underway!":
        action = bot.send_message(message.chat.id, text="Tell me your name, sea wolf  ")
        bot.register_next_step_handler(action, name_handler)

    elif message.text == "Create CV":
        global doc
        doc = Document("cv_form.docx")
        action = bot.send_message(
            message.chat.id, text="Enter your surname and your forename: "
        )
        bot.register_next_step_handler(action, available_date)

    elif message.text == "Add visa":
        keyboard = types.ReplyKeyboardMarkup(True, True)
        keyboard.add("Add visa")
        keyboard.add("Continue with CV")

    elif message.text == "Add endorsement":
        keyboard = types.ReplyKeyboardMarkup(True, True)
        keyboard.add("Add endorsement")
        keyboard.add("Add endorsement")

    elif message.text == "Send CV":
        pass

    elif message.text == "Add certificate":
        pass

    elif message.text == "Add previous experience":
        pass


def name_handler(message):
    """

    :param message:
    :return:
    """
    choice_menu = types.ReplyKeyboardMarkup(True, False)
    choice_menu.row("Create CV")
    choice_menu.row("Send CV(beta)")
    bot.send_message(
        message.chat.id,
        f"Nice to meet you, {message.text} 👋! What would you like to do?",
        reply_markup=choice_menu,
    )


def available_date(message):
    """

    :param message:
    :return:
    """
    sure_and_forename = message.text.split()
    doc.tables[1].column_cells(1)[1].text += sure_and_forename[0]
    doc.tables[1].column_cells(0)[1].text += sure_and_forename[1]
    ready = bot.send_message(
        message.chat.id, text="When are you ready to join a vessel❓"
    )
    bot.register_next_step_handler(ready, rank)


def rank(message):
    doc.tables[0].column_cells(1)[1].text += message.text

    seafarer_rank = bot.send_message(message.chat.id, text="Enter preferred rank: ")
    bot.register_next_step_handler(seafarer_rank, vessel)


def vessel(message):
    doc.tables[0].column_cells(0)[1].text += message.text
    vessel_type = bot.send_message(
        message.chat.id,
        text="Enter preferred type of vessel 🚢",
    )
    bot.register_next_step_handler(vessel_type, date)


def date(message):
    """

    :param message:
    :return:
    """
    doc.tables[0].column_cells(1)[2].text += message.text
    user_name = bot.send_message(message.chat.id, text="Enter your date of birth ")
    bot.register_next_step_handler(user_name, place)


def place(message):
    """

    :param message:
    :return:
    """
    doc.tables[1].column_cells(0)[2].text += message.text
    send = bot.send_message(message.chat.id, text="Enter your place of birth👶")
    bot.register_next_step_handler(send, phone)


def phone(message):
    """

    :param message:
    :return:
    """
    doc.tables[1].column_cells(1)[2].text += message.text
    keyboard = types.ReplyKeyboardMarkup(True, False)
    button_share_phone = types.KeyboardButton(
        text="Share your phone number", request_contact=True
    )
    keyboard.add(button_share_phone)
    number = bot.send_message(
        message.chat.id,
        text="Enter your phone number",
        reply_markup=keyboard,
    )
    bot.register_next_step_handler(number, email)


def email(message):
    doc.tables[1].column_cells(0)[4].text += message.text
    e_box = bot.send_message(message.chat.id, text="Enter your email: ")
    bot.register_next_step_handler(e_box, nationality)


def nationality(message):
    """

    :param message:
    :return:
    """
    # doc.tables[1].column_cells(0)[4].text += message.text
    nation = bot.send_message(message.chat.id, text="Enter your nationality: ")
    bot.register_next_step_handler(nation, home_address)


def home_address(message):
    """

    :param message:
    :return:
    """
    doc.tables[1].column_cells(2)[2].text += message.text
    address = bot.send_message(message.chat.id, text="Enter your home address: 🏠 ")
    bot.register_next_step_handler(address, local_airport)


def local_airport(message):
    """

    :param message:
    :return:
    """
    doc.tables[1].column_cells(0)[3].text += message.text
    air = bot.send_message(
        message.chat.id,
        text="""Enter nearest airport to your location ✈""",
    )
    bot.register_next_step_handler(air, international_airport)


def international_airport(message):
    """

    :param message:
    :return:
    """
    doc.tables[2].column_cells(0)[0].text += message.text
    air = bot.send_message(
        message.chat.id,
        text="""Enter nearest international airport to your location  ✈""",
    )
    bot.register_next_step_handler(air, passport)


def passport(message):
    """

    :param message:
    :return:
    """
    doc.tables[2].column_cells(1)[0].text += message.text
    pasport = bot.send_message(
        message.chat.id,
        text="Enter your passport details: 🛂",
    )
    bot.register_next_step_handler(pasport, seabook)


def seabook(message):
    """

    :param message:
    :return:
    """
    doc.tables[4].column_cells(0)[0].text += message.text
    sea = bot.send_message(
        message.chat.id,
        text="Enter your seabook details ",
    )
    bot.register_next_step_handler(sea, visa)
    doc.save("demo.docx")


def visa(message):
    """

    :param message:
    :return:
    """
    keyboard = types.ReplyKeyboardMarkup(True, True)
    keyboard.add("Добавить визу")
    bot.send_message(
        message.chat.id, text="Add information about your visas: "
    )


def nation_licence(message):
    """

    :param message:
    :return:
    """
    licence = bot.send_message(
        message.chat.id,
        text="""Enter information about your national license """,
    )
    bot.register_next_step_handler(licence, endorsement)


def endorsement(message):
    """

    :param message:
    :return:
    """
    keyboard = types.ReplyKeyboardMarkup(True, True)
    keyboard.add("Add endorsement")
    if_visa = bot.send_message(
        message.chat.id, text="Enter list of your endorsements:  "
    )


bot.polling()
